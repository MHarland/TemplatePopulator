import docx
from . import cfg
import subprocess
import os
from uuid import uuid4
import shutil


class Document:
    def __init__(self, local_template_path: str) -> None:
        self._doc = docx.Document(local_template_path)

    def save_as_pdf(self, path: str):
        """Save document as pdf. Calls Libre Office in a subprocess.

        Args:
            path (str): Path to pdf, relative to WORK_DIR

        Raises:
            Exception: Subprocess returns 1
        """
        tmp_docx_path = os.path.join(cfg.WORK_DIR, f"{uuid4().hex}.docx")
        out_dir = os.path.dirname(path)
        os.makedirs(cfg.WORK_DIR, exist_ok=True)
        self._doc.save(tmp_docx_path)
        command = (
            f"{cfg.SOFFICE}"
            f" --convert-to pdf"
            f" --outdir {out_dir}"
            f" {tmp_docx_path}"
        ).split(" ")
        out = subprocess.run(command, capture_output=True)
        if out.returncode == 1:
            print(out.stderr.decode())
            print(out.stdout.decode())
            raise Exception("Error during loffice execution to save as pdf")
        tmp_pdf_path = tmp_docx_path.rsplit(".")[0] + ".pdf"
        shutil.move(tmp_pdf_path, path)
        os.remove(tmp_docx_path)

    def map_placeholders(self, placeholder_map: dict[str, str]):
        for paragraph in self._doc.paragraphs:
            text = paragraph.text
            for placeholder, target in placeholder_map.items():
                text = text.replace(placeholder, target)
            paragraph.text = text

    def get_whole_text(self) -> str:
        whole_text = ""
        for paragraph in self._doc.paragraphs:
            whole_text += paragraph.text
        return whole_text
